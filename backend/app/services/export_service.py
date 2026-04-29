"""Professional DXC-style export builders for PDF and DOCX delivery reports."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


DXC_NAVY = colors.HexColor("#0F2747")
DXC_BLUE = colors.HexColor("#1D4F91")
DXC_LIGHT_BLUE = colors.HexColor("#EAF1F8")
DXC_GRAY = colors.HexColor("#667085")
DXC_BORDER = colors.HexColor("#D0D5DD")
DXC_SOFT = colors.HexColor("#F5F7FA")
DXC_TEXT = colors.HexColor("#101828")
DXC_WHITE = colors.white
CONFIDENTIALITY_LABEL = "Confidential - DXC Technology Morocco"


@dataclass
class ReportSolution:
    """Normalized view of one exported solution across all report sections."""

    id: str
    name: str
    relevance: float
    overall: float
    description: str
    source: str
    features: list[str]
    business_impact: str
    maturity_level: str
    gap_analysis: dict[str, Any]
    recommendation: dict[str, Any]


def _safe_text(value: Any, default: str = "Data unavailable") -> str:
    """Return a printable string without inventing missing data."""
    if value is None:
        return default
    text = str(value).strip()
    return text or default


def _safe_list(value: Any) -> list[str]:
    """Normalize any iterable-ish payload field into a list of short strings."""
    if not isinstance(value, list):
        return []
    normalized: list[str] = []
    for item in value:
        if isinstance(item, dict):
            title = _safe_text(item.get("title"), "").strip()
            name = _safe_text(item.get("name"), "").strip()
            action = _safe_text(item.get("action"), "").strip()
            description = _safe_text(item.get("description"), "").strip()
            related_feature = item.get("related_feature_missing")
            if isinstance(related_feature, list):
                related_feature = ", ".join(str(entry).strip() for entry in related_feature if str(entry).strip())
            related_resource = _safe_text(item.get("related_resource_needed"), "").strip()

            summary = title or name or action or description
            if title and isinstance(related_feature, str) and related_feature.strip():
                summary = f"{title} - covers {related_feature.strip()}"
            elif title and related_resource:
                summary = f"{title} - resource {related_resource}"
            if summary:
                normalized.append(summary)
            continue

        text = str(item).strip()
        if text:
            normalized.append(text)
    return normalized


def _utc_timestamp() -> str:
    """Build a stable UTC generation timestamp."""
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _horizon_label(horizon: str | None) -> str:
    """Translate horizon codes into presentation labels."""
    labels = {
        "court_terme": "Short term",
        "moyen_terme": "Mid term",
        "long_terme": "Long term",
    }
    return labels.get(str(horizon or "").strip(), "Data unavailable")


def _status_label(status: str | None) -> str:
    """Translate workflow statuses into user-facing labels."""
    labels = {
        "draft": "Draft",
        "submitted": "Submitted",
        "in_qualification": "Qualification",
        "in_selection": "Selection",
        "delivery": "Delivery",
        "export_ready": "Export Ready",
        "abandoned": "Abandoned",
    }
    return labels.get(str(status or "").strip(), "Unknown")


def _objective_label(value: str | None) -> str:
    """Translate objective tags into presentation labels."""
    labels = {
        "cost_reduction": "Cost reduction",
        "cx_improvement": "Customer experience improvement",
        "risk_mitigation": "Risk mitigation",
        "market_opportunity": "Market opportunity",
    }
    return labels.get(str(value or "").strip(), _safe_text(value))


def _infer_stage_gates(status: str, rework_note: str | None) -> list[dict[str, str | None]]:
    """Build a fallback stage gate summary when the frontend did not send one."""
    reached = {
        "SG-1": status in {"submitted", "in_qualification", "in_selection", "delivery", "export_ready"},
        "SG-2": status in {"in_qualification", "in_selection", "delivery", "export_ready"},
        "SG-3": status in {"in_selection", "delivery", "export_ready"},
        "SG-4": status == "export_ready",
    }
    comments = {
        "SG-1": rework_note if status == "submitted" else None,
        "SG-2": rework_note if status == "in_qualification" else None,
        "SG-3": rework_note if status == "in_selection" else None,
        "SG-4": rework_note if status in {"delivery", "export_ready"} else None,
    }
    return [
        {
            "gate": "SG-1",
            "phase": "Sourcing",
            "decision": "GO" if reached["SG-1"] else "PENDING",
            "status_after": "submitted" if reached["SG-1"] else "draft",
            "comment": comments["SG-1"],
        },
        {
            "gate": "SG-2",
            "phase": "Discovery",
            "decision": "GO" if reached["SG-2"] else "PENDING",
            "status_after": "in_qualification" if reached["SG-2"] else "submitted",
            "comment": comments["SG-2"],
        },
        {
            "gate": "SG-3",
            "phase": "Qualification",
            "decision": "GO" if reached["SG-3"] else "PENDING",
            "status_after": "in_selection" if reached["SG-3"] else "in_qualification",
            "comment": comments["SG-3"],
        },
        {
            "gate": "SG-4",
            "phase": "Delivery",
            "decision": "GO" if reached["SG-4"] else "PENDING",
            "status_after": "export_ready" if reached["SG-4"] else "delivery",
            "comment": comments["SG-4"],
        },
    ]


def _normalize_solutions(
    selected_solutions: list[dict[str, Any]],
    delivery_solutions: list[dict[str, Any]],
    recommendations: list[dict[str, Any]],
) -> list[ReportSolution]:
    """Merge solution snapshots, delivery metrics, and recommendations into one view."""
    selected_by_id = {str(item.get("id", "")): item for item in selected_solutions if item.get("id")}
    delivery_by_id = {str(item.get("id", "")): item for item in delivery_solutions if item.get("id")}
    recommendation_by_id = {str(item.get("solution_id", "")): item for item in recommendations if item.get("solution_id")}
    recommendation_by_name = {str(item.get("solution_name", "")): item for item in recommendations if item.get("solution_name")}

    merged: list[ReportSolution] = []
    for delivery in delivery_solutions:
        solution_id = str(delivery.get("id", "")).strip()
        solution_name = _safe_text(delivery.get("name"), "Unnamed solution")
        selected = selected_by_id.get(solution_id, {})
        recommendation = recommendation_by_id.get(solution_id) or recommendation_by_name.get(solution_name, {})
        merged.append(
            ReportSolution(
                id=solution_id,
                name=solution_name,
                relevance=float(delivery.get("relevance", 0) or 0),
                overall=float(delivery.get("overall", 0) or 0),
                description=_safe_text(selected.get("description"), "Data unavailable"),
                source=_safe_text(selected.get("source"), "Data unavailable"),
                features=_safe_list(selected.get("features")),
                business_impact=_safe_text(selected.get("business_impact"), "Data unavailable"),
                maturity_level=_safe_text(selected.get("maturity_level"), "Data unavailable"),
                gap_analysis=selected.get("gap_analysis") or {},
                recommendation=recommendation,
            )
        )

    if not merged:
        for selected in selected_solutions:
            solution_id = str(selected.get("id", "")).strip()
            solution_name = _safe_text(selected.get("name"), "Unnamed solution")
            recommendation = recommendation_by_id.get(solution_id) or recommendation_by_name.get(solution_name, {})
            merged.append(
                ReportSolution(
                    id=solution_id,
                    name=solution_name,
                    relevance=float(selected.get("relevance", 0) or 0),
                    overall=float(selected.get("overall", 0) or 0),
                    description=_safe_text(selected.get("description"), "Data unavailable"),
                    source=_safe_text(selected.get("source"), "Data unavailable"),
                    features=_safe_list(selected.get("features")),
                    business_impact=_safe_text(selected.get("business_impact"), "Data unavailable"),
                    maturity_level=_safe_text(selected.get("maturity_level"), "Data unavailable"),
                    gap_analysis=selected.get("gap_analysis") or {},
                    recommendation=recommendation,
                )
            )

    return merged


def _build_executive_summary(solutions: list[ReportSolution]) -> str:
    """Create a short consulting-style executive summary."""
    if not solutions:
        return (
            "No delivery solution has been selected yet. The report therefore captures the business need "
            "and the current delivery recommendations without a validated solution package."
        )

    best = max(solutions, key=lambda item: item.overall)
    ivi_scores = [float(item.gap_analysis.get("ivi_score", 0) or 0) for item in solutions if item.gap_analysis]
    fit_scores = [int(item.gap_analysis.get("fit_score", 0) or 0) for item in solutions if item.gap_analysis]
    average_ivi = sum(ivi_scores) / len(ivi_scores) if ivi_scores else None
    average_fit = sum(fit_scores) / len(fit_scores) if fit_scores else None

    summary = [
        f"{len(solutions)} delivery solution(s) were retained for the final recommendation set.",
        f"The leading option is {best.name} with an overall delivery score of {best.overall:.2f} and a discovery relevance of {best.relevance:.0f}%.",
    ]
    if average_ivi is not None:
        summary.append(f"The average IVI readiness across retained solutions is {average_ivi:.1f}/100.")
    if average_fit is not None:
        summary.append(f"The average gap fit score is {average_fit:.1f}/10.")
    return " ".join(summary)


def _build_solution_roadmap(solution: ReportSolution) -> list[tuple[str, str, str]]:
    """Build a simple J+30 / J+60 / J+90 roadmap from existing solution evidence."""
    technical = _safe_list(solution.recommendation.get("technical_recommendations"))
    organizational = _safe_list(solution.recommendation.get("organizational_recommendations"))
    missing = _safe_list(solution.gap_analysis.get("features_missing"))

    return [
        (
            "J+30",
            "Mobilization and prerequisites",
            "; ".join((technical[:2] + organizational[:1] + missing[:1])[:3]) or "Confirm governance, scope, and delivery prerequisites.",
        ),
        (
            "J+60",
            "Build and integration",
            "; ".join((technical[2:4] + organizational[1:3] + missing[1:2])[:3]) or "Implement priority integrations, data flows, and pilot capabilities.",
        ),
        (
            "J+90",
            "Industrialization and value tracking",
            "; ".join((technical[4:6] + organizational[3:5])[:3]) or "Stabilize operations, enable adoption, and track KPI uplift.",
        ),
    ]


def _build_risk_rows(solution: ReportSolution) -> list[tuple[str, str]]:
    """Pair each risk with a practical mitigation drawn from available recommendations."""
    risks = _safe_list(solution.gap_analysis.get("risks"))
    technical = _safe_list(solution.recommendation.get("technical_recommendations"))
    organizational = _safe_list(solution.recommendation.get("organizational_recommendations"))
    mitigations = technical + organizational

    rows: list[tuple[str, str]] = []
    for index, risk in enumerate(risks):
        mitigation = mitigations[index] if index < len(mitigations) else "Assign an owner, define a mitigation checkpoint, and monitor until closure."
        rows.append((risk, mitigation))

    if not rows:
        rows.append(("No explicit delivery risk provided", "Continue monitoring scope, dependencies, and data readiness during delivery governance."))
    return rows


def _build_business_need_metadata(tags: dict[str, Any], horizon: str, status: str) -> list[list[str]]:
    """Create key-value rows for the business need summary box."""
    domains = ", ".join(_safe_list(tags.get("domaine"))) or "Data unavailable"
    impacts = ", ".join(_safe_list(tags.get("impact"))) or "Data unavailable"
    return [
        ["Project", "IPM Flow"],
        ["Organization", "DXC Technology Morocco"],
        ["Workflow status", _status_label(status)],
        ["Horizon", _horizon_label(horizon)],
        ["Primary objective", _objective_label(tags.get("objectif"))],
        ["Domains", domains],
        ["Impacts", impacts],
        ["Origin", _safe_text(tags.get("origine"))],
    ]


def _build_pdf_styles() -> dict[str, ParagraphStyle]:
    """Centralize PDF typography and spacing."""
    sample = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title",
            parent=sample["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=DXC_NAVY,
            spaceAfter=10,
        ),
        "cover_meta": ParagraphStyle(
            "cover_meta",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=DXC_GRAY,
            spaceAfter=8,
        ),
        "section": ParagraphStyle(
            "section",
            parent=sample["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=DXC_NAVY,
            spaceBefore=10,
            spaceAfter=10,
        ),
        "subsection": ParagraphStyle(
            "subsection",
            parent=sample["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=DXC_BLUE,
            spaceBefore=8,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "body",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=14,
            textColor=DXC_TEXT,
        ),
        "small": ParagraphStyle(
            "small",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=DXC_GRAY,
        ),
        "badge": ParagraphStyle(
            "badge",
            parent=sample["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            textColor=DXC_WHITE,
            alignment=1,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=sample["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            leftIndent=10,
            bulletIndent=0,
            textColor=DXC_TEXT,
        ),
    }


def _pdf_boxed_paragraph(text: str, style: ParagraphStyle, width: float) -> Table:
    """Render a single boxed paragraph block."""
    table = Table([[Paragraph(text, style)]], colWidths=[width])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), DXC_SOFT),
                ("BOX", (0, 0), (-1, -1), 0.75, DXC_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ]
        )
    )
    return table


def _pdf_key_value_table(rows: list[list[str]], width: float) -> Table:
    """Render a structured key-value table."""
    label_width = width * 0.28
    value_width = width - label_width
    table = Table(rows, colWidths=[label_width, value_width], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), DXC_LIGHT_BLUE),
                ("BACKGROUND", (1, 0), (1, -1), DXC_WHITE),
                ("TEXTCOLOR", (0, 0), (-1, -1), DXC_TEXT),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, DXC_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _pdf_standard_table(rows: list[list[str]], widths: list[float], header_rows: int = 1) -> Table:
    """Render a clean DXC table."""
    table = Table(rows, colWidths=widths, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, header_rows - 1), DXC_NAVY),
                ("TEXTCOLOR", (0, 0), (-1, header_rows - 1), DXC_WHITE),
                ("FONTNAME", (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
                ("FONTNAME", (0, header_rows), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.8),
                ("GRID", (0, 0), (-1, -1), 0.5, DXC_BORDER),
                ("ROWBACKGROUNDS", (0, header_rows), (-1, -1), [DXC_WHITE, DXC_SOFT]),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _pdf_bullets(items: list[str], styles: dict[str, ParagraphStyle]) -> list[Paragraph]:
    """Render bullet paragraphs with stable spacing."""
    if not items:
        return [Paragraph("No item provided.", styles["body"])]
    return [Paragraph(_safe_text(item), styles["bullet"], bulletText="-") for item in items]


def _draw_pdf_header_footer(canvas, doc) -> None:
    """Draw DXC header/footer with page number on every page."""
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 9)
    canvas.setFillColor(DXC_NAVY)
    canvas.drawString(doc.leftMargin, A4[1] - 14 * mm, "DXC Technology | IPM Flow")
    canvas.setStrokeColor(DXC_BORDER)
    canvas.setLineWidth(0.5)
    canvas.line(doc.leftMargin, A4[1] - 16 * mm, A4[0] - doc.rightMargin, A4[1] - 16 * mm)

    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(DXC_GRAY)
    footer_y = 11 * mm
    canvas.line(doc.leftMargin, footer_y + 5, A4[0] - doc.rightMargin, footer_y + 5)
    canvas.drawString(doc.leftMargin, footer_y - 2, CONFIDENTIALITY_LABEL)
    canvas.drawRightString(A4[0] - doc.rightMargin, footer_y - 2, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def build_pdf_report(
    *,
    need_id: str,
    pitch: str,
    horizon: str,
    tags: dict[str, Any],
    status: str,
    rework_note: str | None,
    recommendations: list[dict[str, Any]],
    delivery_solutions: list[dict[str, Any]],
    selected_solutions: list[dict[str, Any]],
    stage_gates: list[dict[str, Any]],
) -> bytes:
    """Build a polished PDF report while keeping the export endpoint stable."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=22 * mm,
        bottomMargin=18 * mm,
        title=f"Innovation Delivery Recommendations - {need_id}",
        author="IPM Flow",
    )

    styles = _build_pdf_styles()
    content_width = A4[0] - doc.leftMargin - doc.rightMargin
    generated_at = _utc_timestamp()
    solutions = _normalize_solutions(selected_solutions, delivery_solutions, recommendations)
    stage_gate_rows = stage_gates or _infer_stage_gates(status, rework_note)

    story: list[Any] = []

    cover_band = Table([[""]], colWidths=[content_width], rowHeights=[18 * mm])
    cover_band.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), DXC_NAVY)]))
    story.extend(
        [
            Spacer(1, 14 * mm),
            cover_band,
            Spacer(1, 18 * mm),
            Paragraph("Innovation Delivery Recommendations", styles["title"]),
            Paragraph("Consulting-style delivery recommendation report", styles["cover_meta"]),
            Spacer(1, 10 * mm),
            _pdf_key_value_table(
                [
                    ["Need ID", need_id],
                    ["Generated on", generated_at],
                    ["Project", "IPM Flow"],
                    ["Entity", "DXC Technology Morocco"],
                    ["Confidentiality", CONFIDENTIALITY_LABEL],
                ],
                content_width,
            ),
            Spacer(1, 12 * mm),
            _pdf_boxed_paragraph(
                "This document summarizes the current business need, the retained DXC-aligned delivery options, "
                "their IVI qualification signals, and the operational recommendations required to move toward delivery.",
                styles["body"],
                content_width,
            ),
            PageBreak(),
        ]
    )

    story.append(Paragraph("Executive Summary", styles["section"]))
    story.append(_pdf_boxed_paragraph(_build_executive_summary(solutions), styles["body"], content_width))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Business Need", styles["section"]))
    story.append(_pdf_boxed_paragraph(_safe_text(pitch, "Data unavailable"), styles["body"], content_width))
    story.append(Spacer(1, 8))
    story.append(_pdf_key_value_table(_build_business_need_metadata(tags, horizon, status), content_width))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Selected Solutions", styles["section"]))
    selected_rows = [["Solution", "Source", "Relevance", "Overall"]]
    for solution in solutions:
        selected_rows.append([
            solution.name,
            solution.source,
            f"{solution.relevance:.0f}%",
            f"{solution.overall:.2f}",
        ])
    if len(selected_rows) == 1:
        selected_rows.append(["No selected solution", "Data unavailable", "0%", "0.00"])
    story.append(_pdf_standard_table(selected_rows, [content_width * 0.34, content_width * 0.34, content_width * 0.14, content_width * 0.18]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("IVI Score Summary", styles["section"]))
    if solutions:
        for solution in solutions:
            gap = solution.gap_analysis or {}
            ivi = gap.get("ivi_scoring") or {}
            badge_row = [[
                Paragraph(f"{solution.name}<br/><font size='9'>IVI {float(gap.get('ivi_score', 0) or 0):.1f}/100</font>", styles["badge"]),
                Paragraph(f"Maturity<br/>{int((ivi.get('maturite') or {}).get('score', 0) or 0)}/5", styles["badge"]),
                Paragraph(f"Expertise<br/>{int((ivi.get('expertise') or {}).get('score', 0) or 0)}/5", styles["badge"]),
                Paragraph(f"Duration<br/>{int((ivi.get('duree') or {}).get('score', 0) or 0)}/5", styles["badge"]),
                Paragraph(f"Data<br/>{int((ivi.get('donnees') or {}).get('score', 0) or 0)}/5", styles["badge"]),
                Paragraph(f"Business impact<br/>{int((ivi.get('impact_business') or {}).get('score', 0) or 0)}/5", styles["badge"]),
            ]]
            badge_table = Table(
                badge_row,
                colWidths=[
                    content_width * 0.24,
                    content_width * 0.152,
                    content_width * 0.152,
                    content_width * 0.152,
                    content_width * 0.152,
                    content_width * 0.152,
                ],
            )
            badge_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, 0), DXC_NAVY),
                        ("BACKGROUND", (1, 0), (-1, 0), DXC_BLUE),
                        ("BOX", (0, 0), (-1, -1), 0.5, DXC_BORDER),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )
            story.append(badge_table)
            story.append(Spacer(1, 6))
            client_message = _safe_text(gap.get("client_message"), "Data unavailable")
            story.append(_pdf_boxed_paragraph(client_message, styles["body"], content_width))
            story.append(Spacer(1, 10))
    else:
        story.append(_pdf_boxed_paragraph("IVI scoring is unavailable because no selected solution payload was provided.", styles["body"], content_width))

    story.append(Paragraph("Gap Analysis", styles["section"]))
    if solutions:
        for solution in solutions:
            gap = solution.gap_analysis or {}
            story.append(Paragraph(solution.name, styles["subsection"]))
            story.append(
                _pdf_key_value_table(
                    [
                        ["Fit score", f"{int(gap.get('fit_score', 0) or 0)}/10"],
                        ["Feasibility", f"{int((gap.get('feasibility') or {}).get('score', 0) or 0)}/5"],
                        ["Maturity level", solution.maturity_level],
                        ["Business impact", solution.business_impact],
                    ],
                    content_width,
                )
            )
            story.append(Spacer(1, 6))
            story.append(_pdf_boxed_paragraph(_safe_text(gap.get("fit_justification")), styles["body"], content_width))
            story.append(Spacer(1, 6))

            matching = _safe_list(gap.get("features_matching"))
            missing = _safe_list(gap.get("features_missing"))
            risks = _safe_list(gap.get("risks"))
            resources = _safe_list(gap.get("resources_needed"))

            gap_rows = [
                [
                    Paragraph("<b>Features matching</b><br/>" + ("<br/>".join(f"- {item}" for item in matching) or "Data unavailable"), styles["body"]),
                    Paragraph("<b>Features missing</b><br/>" + ("<br/>".join(f"- {item}" for item in missing) or "Data unavailable"), styles["body"]),
                ],
                [
                    Paragraph("<b>Risks</b><br/>" + ("<br/>".join(f"- {item}" for item in risks) or "Data unavailable"), styles["body"]),
                    Paragraph("<b>Resources needed</b><br/>" + ("<br/>".join(f"- {item}" for item in resources) or "Data unavailable"), styles["body"]),
                ],
            ]
            gap_table = Table(gap_rows, colWidths=[content_width * 0.5, content_width * 0.5])
            gap_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), DXC_SOFT),
                        ("BOX", (0, 0), (-1, -1), 0.5, DXC_BORDER),
                        ("INNERGRID", (0, 0), (-1, -1), 0.5, DXC_BORDER),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(gap_table)
            story.append(Spacer(1, 10))
    else:
        story.append(_pdf_boxed_paragraph("Gap analysis details are unavailable because no solution snapshot was attached to the export payload.", styles["body"], content_width))

    story.append(Paragraph("Technical Recommendations", styles["section"]))
    for solution in solutions:
        story.append(Paragraph(solution.name, styles["subsection"]))
        story.extend(_pdf_bullets(_safe_list(solution.recommendation.get("technical_recommendations")), styles))
        story.append(Spacer(1, 8))

    story.append(Paragraph("Organizational Recommendations", styles["section"]))
    for solution in solutions:
        story.append(Paragraph(solution.name, styles["subsection"]))
        story.extend(_pdf_bullets(_safe_list(solution.recommendation.get("organizational_recommendations")), styles))
        story.append(Spacer(1, 8))

    story.append(Paragraph("KPIs", styles["section"]))
    kpi_rows = [["Solution", "KPI", "Target", "Measurement criteria"]]
    for solution in solutions:
        kpis = solution.recommendation.get("kpis") if isinstance(solution.recommendation.get("kpis"), list) else []
        if not kpis:
            kpi_rows.append([solution.name, "Data unavailable", "Data unavailable", "Data unavailable"])
            continue
        for kpi in kpis:
            if not isinstance(kpi, dict):
                continue
            kpi_rows.append([
                solution.name,
                _safe_text(kpi.get("name")),
                _safe_text(kpi.get("target")),
                _safe_text(kpi.get("measurement_criteria")),
            ])
    story.append(_pdf_standard_table(kpi_rows, [content_width * 0.22, content_width * 0.2, content_width * 0.24, content_width * 0.34]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Delivery Roadmap J+30 / J+60 / J+90", styles["section"]))
    roadmap_rows = [["Solution", "Milestone", "Focus", "Actions"]]
    for solution in solutions:
        for milestone, focus, actions in _build_solution_roadmap(solution):
            roadmap_rows.append([solution.name, milestone, focus, actions])
    if len(roadmap_rows) == 1:
        roadmap_rows.append(["Data unavailable", "J+30", "Data unavailable", "Data unavailable"])
    story.append(_pdf_standard_table(roadmap_rows, [content_width * 0.19, content_width * 0.11, content_width * 0.22, content_width * 0.48]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Risks and Mitigation", styles["section"]))
    risk_rows = [["Solution", "Risk", "Mitigation"]]
    for solution in solutions:
        for risk, mitigation in _build_risk_rows(solution):
            risk_rows.append([solution.name, risk, mitigation])
    story.append(_pdf_standard_table(risk_rows, [content_width * 0.2, content_width * 0.34, content_width * 0.46]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Stage Gates Validation Summary", styles["section"]))
    gate_rows = [["Gate", "Phase", "Decision", "Status after gate", "Comment"]]
    for gate in stage_gate_rows:
        gate_rows.append([
            _safe_text(gate.get("gate")),
            _safe_text(gate.get("phase")),
            _safe_text(gate.get("decision")),
            _status_label(str(gate.get("status_after") or "")),
            _safe_text(gate.get("comment"), "No comment"),
        ])
    story.append(_pdf_standard_table(gate_rows, [content_width * 0.1, content_width * 0.18, content_width * 0.14, content_width * 0.2, content_width * 0.38]))

    doc.build(story, onFirstPage=_draw_pdf_header_footer, onLaterPages=_draw_pdf_header_footer)
    return buffer.getvalue()


def _set_cell_shading(cell, fill: str) -> None:
    """Apply background color to a DOCX table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_text(cell, text: str, bold: bool = False, color: Any | None = None) -> None:
    """Write styled text into a DOCX cell."""
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def _docx_add_page_number(paragraph) -> None:
    """Insert a PAGE field into a DOCX paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _configure_docx_header_footer(document) -> None:
    """Apply a consistent DXC header/footer to each document section."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        header = section.header
        header.is_linked_to_previous = False
        header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_p.add_run("DXC Technology | IPM Flow")
        header_run.bold = True
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor(0x0F, 0x27, 0x47)

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_p.add_run(f"{CONFIDENTIALITY_LABEL} | Page ")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
        _docx_add_page_number(footer_p)


def _style_docx_heading(paragraph, level: int) -> None:
    """Apply DXC colors to generated headings."""
    from docx.shared import Pt, RGBColor

    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x27, 0x47)
        run.bold = True
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)


def _docx_key_value_table(document, rows: list[list[str]]) -> None:
    """Insert a professional key-value table."""
    from docx.shared import RGBColor

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, color=RGBColor(0x10, 0x18, 0x28))
        _set_cell_shading(cells[0], "EAF1F8")
        _set_cell_text(cells[1], value)


def _docx_standard_table(document, rows: list[list[str]]) -> None:
    """Insert a generic report table with a dark header."""
    from docx.shared import RGBColor

    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        _set_cell_text(header_cells[index], value, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_shading(header_cells[index], "0F2747")

    for data_row in rows[1:]:
        row_cells = table.add_row().cells
        for index, value in enumerate(data_row):
            _set_cell_text(row_cells[index], value)


def build_docx_report(
    *,
    need_id: str,
    pitch: str,
    horizon: str,
    tags: dict[str, Any],
    status: str,
    rework_note: str | None,
    recommendations: list[dict[str, Any]],
    delivery_solutions: list[dict[str, Any]],
    selected_solutions: list[dict[str, Any]],
    stage_gates: list[dict[str, Any]],
) -> bytes:
    """Build a more structured DOCX report with the same export endpoint."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Pt, RGBColor

    document = Document()
    _configure_docx_header_footer(document)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    generated_at = _utc_timestamp()
    solutions = _normalize_solutions(selected_solutions, delivery_solutions, recommendations)
    stage_gate_rows = stage_gates or _infer_stage_gates(status, rework_note)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Innovation Delivery Recommendations")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x0F, 0x27, 0x47)

    subtitle = document.add_paragraph("Consulting-style delivery recommendation report")
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x70, 0x85)

    _docx_key_value_table(
        document,
        [
            ["Need ID", need_id],
            ["Generated on", generated_at],
            ["Project", "IPM Flow"],
            ["Entity", "DXC Technology Morocco"],
            ["Confidentiality", CONFIDENTIALITY_LABEL],
        ],
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    heading = document.add_heading("Executive Summary", level=2)
    _style_docx_heading(heading, 2)
    document.add_paragraph(_build_executive_summary(solutions))

    heading = document.add_heading("Business Need", level=2)
    _style_docx_heading(heading, 2)
    document.add_paragraph(_safe_text(pitch, "Data unavailable"))
    _docx_key_value_table(document, _build_business_need_metadata(tags, horizon, status))

    heading = document.add_heading("Selected Solutions", level=2)
    _style_docx_heading(heading, 2)
    selected_rows = [["Solution", "Source", "Relevance", "Overall"]]
    for solution in solutions:
        selected_rows.append([
            solution.name,
            solution.source,
            f"{solution.relevance:.0f}%",
            f"{solution.overall:.2f}",
        ])
    if len(selected_rows) == 1:
        selected_rows.append(["No selected solution", "Data unavailable", "0%", "0.00"])
    _docx_standard_table(document, selected_rows)

    heading = document.add_heading("IVI Score Summary", level=2)
    _style_docx_heading(heading, 2)
    ivi_rows = [["Solution", "IVI", "Maturity", "Expertise", "Duration", "Data", "Business impact"]]
    for solution in solutions:
        ivi = (solution.gap_analysis or {}).get("ivi_scoring") or {}
        ivi_rows.append([
            solution.name,
            f"{float((solution.gap_analysis or {}).get('ivi_score', 0) or 0):.1f}/100",
            f"{int((ivi.get('maturite') or {}).get('score', 0) or 0)}/5",
            f"{int((ivi.get('expertise') or {}).get('score', 0) or 0)}/5",
            f"{int((ivi.get('duree') or {}).get('score', 0) or 0)}/5",
            f"{int((ivi.get('donnees') or {}).get('score', 0) or 0)}/5",
            f"{int((ivi.get('impact_business') or {}).get('score', 0) or 0)}/5",
        ])
    if len(ivi_rows) == 1:
        ivi_rows.append(["Data unavailable", "0/100", "0/5", "0/5", "0/5", "0/5", "0/5"])
    _docx_standard_table(document, ivi_rows)

    heading = document.add_heading("Gap Analysis", level=2)
    _style_docx_heading(heading, 2)
    for solution in solutions:
        subheading = document.add_heading(solution.name, level=3)
        _style_docx_heading(subheading, 3)
        _docx_key_value_table(
            document,
            [
                ["Fit score", f"{int((solution.gap_analysis or {}).get('fit_score', 0) or 0)}/10"],
                ["Feasibility", f"{int(((solution.gap_analysis or {}).get('feasibility') or {}).get('score', 0) or 0)}/5"],
                ["Maturity level", solution.maturity_level],
                ["Business impact", solution.business_impact],
            ],
        )
        document.add_paragraph(_safe_text((solution.gap_analysis or {}).get("fit_justification")))
        document.add_paragraph("Features matching").runs[0].bold = True
        for item in _safe_list((solution.gap_analysis or {}).get("features_matching")) or ["Data unavailable"]:
            document.add_paragraph(item, style="List Bullet")
        document.add_paragraph("Features missing").runs[0].bold = True
        for item in _safe_list((solution.gap_analysis or {}).get("features_missing")) or ["Data unavailable"]:
            document.add_paragraph(item, style="List Bullet")
        document.add_paragraph("Risks").runs[0].bold = True
        for item in _safe_list((solution.gap_analysis or {}).get("risks")) or ["Data unavailable"]:
            document.add_paragraph(item, style="List Bullet")

    heading = document.add_heading("Technical Recommendations", level=2)
    _style_docx_heading(heading, 2)
    for solution in solutions:
        subheading = document.add_heading(solution.name, level=3)
        _style_docx_heading(subheading, 3)
        for item in _safe_list(solution.recommendation.get("technical_recommendations")) or ["Data unavailable"]:
            document.add_paragraph(item, style="List Bullet")

    heading = document.add_heading("Organizational Recommendations", level=2)
    _style_docx_heading(heading, 2)
    for solution in solutions:
        subheading = document.add_heading(solution.name, level=3)
        _style_docx_heading(subheading, 3)
        for item in _safe_list(solution.recommendation.get("organizational_recommendations")) or ["Data unavailable"]:
            document.add_paragraph(item, style="List Bullet")

    heading = document.add_heading("KPIs", level=2)
    _style_docx_heading(heading, 2)
    kpi_rows = [["Solution", "KPI", "Target", "Measurement criteria"]]
    for solution in solutions:
        kpis = solution.recommendation.get("kpis") if isinstance(solution.recommendation.get("kpis"), list) else []
        if not kpis:
            kpi_rows.append([solution.name, "Data unavailable", "Data unavailable", "Data unavailable"])
            continue
        for kpi in kpis:
            if not isinstance(kpi, dict):
                continue
            kpi_rows.append([
                solution.name,
                _safe_text(kpi.get("name")),
                _safe_text(kpi.get("target")),
                _safe_text(kpi.get("measurement_criteria")),
            ])
    _docx_standard_table(document, kpi_rows)

    heading = document.add_heading("Delivery Roadmap J+30 / J+60 / J+90", level=2)
    _style_docx_heading(heading, 2)
    roadmap_rows = [["Solution", "Milestone", "Focus", "Actions"]]
    for solution in solutions:
        for milestone, focus, actions in _build_solution_roadmap(solution):
            roadmap_rows.append([solution.name, milestone, focus, actions])
    _docx_standard_table(document, roadmap_rows)

    heading = document.add_heading("Risks and Mitigation", level=2)
    _style_docx_heading(heading, 2)
    risk_rows = [["Solution", "Risk", "Mitigation"]]
    for solution in solutions:
        for risk, mitigation in _build_risk_rows(solution):
            risk_rows.append([solution.name, risk, mitigation])
    _docx_standard_table(document, risk_rows)

    heading = document.add_heading("Stage Gates Validation Summary", level=2)
    _style_docx_heading(heading, 2)
    gate_rows = [["Gate", "Phase", "Decision", "Status after gate", "Comment"]]
    for gate in stage_gate_rows:
        gate_rows.append([
            _safe_text(gate.get("gate")),
            _safe_text(gate.get("phase")),
            _safe_text(gate.get("decision")),
            _status_label(str(gate.get("status_after") or "")),
            _safe_text(gate.get("comment"), "No comment"),
        ])
    _docx_standard_table(document, gate_rows)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
